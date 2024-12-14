import os
from docx import Document

# Hàm lấy cấu trúc thư mục
def list_files(startpath):
    file_list = []
    for root, dirs, files in os.walk(startpath):
        for file in files:
            file_list.append(os.path.join(root, file))
        for dir in dirs:
            file_list.append(os.path.join(root, dir))
    return file_list

# Hàm tạo cấu trúc thư mục dưới dạng cây
def display_tree(startpath, file, prefix="", line_count=0):
    # Viết tên thư mục hoặc tệp vào tài liệu
    file.write(prefix + os.path.basename(startpath) + "\n")
    line_count += 1

    # Nếu đã đạt đến 30 dòng, thêm một dòng trống
    if line_count >= 50:
        file.write("\n")
        line_count = 0  # Reset dòng đếm

    prefix += "├── "
    if os.path.isdir(startpath):
        for item in os.listdir(startpath):
            item_path = os.path.join(startpath, item)
            if os.path.isdir(item_path):
                line_count = display_tree(item_path, file, prefix, line_count)
            else:
                file.write(prefix + item + "\n")
                line_count += 1

                # Thêm dòng trống nếu đã đạt 30 dòng
                if line_count >= 30:
                    file.write("\n")
                    line_count = 0  # Reset dòng đếm
    return line_count

# Tạo tài liệu Word
doc = Document()
doc.add_heading('Cấu trúc Thư mục Dự án', 0)

# Đường dẫn thư mục của bạn
project_folder = 'D:\project\royal_healthy_care\royal-health-care\src\assets\images\drinks'

# Lưu cấu trúc thư mục vào tài liệu Word
with open("project_structure.txt", "w") as f:
    display_tree(project_folder, f)

# Đọc tệp cấu trúc và thêm vào tài liệu Word
with open("project_structure.txt", "r") as f:
    content = f.read()
    doc.add_paragraph(content)

# Lưu tài liệu Word
doc.save("project_structure.docx")
